"""
Document Parser
===============
Robust text extraction from uploaded files using a tiered strategy:

  1. unstructured  — best for complex layouts, tables, multi-column PDFs
  2. pdfplumber    — fallback for formatted PDFs
  3. pypdf         — lightweight fallback for simple PDFs
  4. python-docx   — DOCX with table support
  5. pytesseract   — optional OCR for scanned/image-based PDFs (requires Tesseract binary)

Each tier is wrapped in a try/except so missing optional dependencies fail
gracefully rather than crashing the endpoint.
"""

from __future__ import annotations

import io
import structlog

logger = structlog.get_logger(__name__)

# ─── Result dataclass ─────────────────────────────────────────────────────


class ExtractionResult:
    """Encapsulates the result of a text extraction attempt."""

    __slots__ = ("text", "method", "warning", "is_ocr", "word_count")

    def __init__(
        self,
        text: str,
        method: str,
        warning: str | None = None,
        is_ocr: bool = False,
    ):
        self.text = text.strip()
        self.method = method
        self.warning = warning
        self.is_ocr = is_ocr
        self.word_count = len(self.text.split()) if self.text else 0


# ─── PDF extraction ───────────────────────────────────────────────────────


def _extract_pdf_unstructured(file_bytes: bytes) -> str:
    """Extract PDF text via unstructured (handles complex layouts)."""
    from unstructured.partition.pdf import partition_pdf  # noqa: PLC0415

    elements = partition_pdf(file=io.BytesIO(file_bytes), strategy="fast")
    return "\n".join(str(el) for el in elements if str(el).strip())


def _extract_pdf_pdfplumber(file_bytes: bytes) -> str:
    """Extract PDF text via pdfplumber (good for tables and formatted docs)."""
    import pdfplumber  # noqa: PLC0415

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        pages = [page.extract_text() or "" for page in pdf.pages]
    return "\n".join(pages)


def _extract_pdf_pypdf(file_bytes: bytes) -> str:
    """Extract PDF text via pypdf (lightweight, works on simple PDFs)."""
    from pypdf import PdfReader  # noqa: PLC0415

    reader = PdfReader(io.BytesIO(file_bytes))
    pages = [(page.extract_text() or "") for page in reader.pages]
    return "\n".join(pages)


def _extract_pdf_ocr(file_bytes: bytes) -> str:
    """OCR fallback using pytesseract + pdf2image (requires Tesseract binary)."""
    import pytesseract  # noqa: PLC0415
    from pdf2image import convert_from_bytes  # noqa: PLC0415

    images = convert_from_bytes(file_bytes)
    return "\n".join(pytesseract.image_to_string(img) for img in images)


def extract_text_from_pdf(file_bytes: bytes, use_ocr: bool = False) -> ExtractionResult:
    """
    Extract text from PDF bytes using a tiered strategy.

    Falls through to the next strategy if the current one produces no text
    or raises an exception.
    """
    strategies = [
        ("unstructured", _extract_pdf_unstructured),
        ("pdfplumber", _extract_pdf_pdfplumber),
        ("pypdf", _extract_pdf_pypdf),
    ]

    for method_name, extractor in strategies:
        try:
            text = extractor(file_bytes)
            if text.strip():
                logger.info("pdf_extraction_success", method=method_name)
                return ExtractionResult(text=text, method=method_name)
            logger.debug("pdf_extraction_empty", method=method_name)
        except Exception as exc:
            logger.warning("pdf_extraction_failed", method=method_name, error=str(exc))

    # All text-based strategies exhausted — try OCR if requested
    if use_ocr:
        try:
            text = _extract_pdf_ocr(file_bytes)
            if text.strip():
                logger.info("pdf_extraction_success", method="ocr")
                return ExtractionResult(text=text, method="ocr", is_ocr=True)
        except Exception as exc:
            logger.warning("pdf_ocr_failed", error=str(exc))

    warning = (
        "Could not extract text from this PDF. "
        "It may be image-based or encrypted. "
        "Try enabling OCR or pasting the text directly."
    )
    logger.warning("pdf_extraction_all_failed")
    return ExtractionResult(text="", method="failed", warning=warning)


# ─── DOCX extraction ──────────────────────────────────────────────────────


def _extract_docx_unstructured(file_bytes: bytes) -> str:
    """Extract DOCX text via unstructured."""
    from unstructured.partition.docx import partition_docx  # noqa: PLC0415

    elements = partition_docx(file=io.BytesIO(file_bytes))
    return "\n".join(str(el) for el in elements if str(el).strip())


def _extract_docx_python_docx(file_bytes: bytes) -> str:
    """Extract DOCX text via python-docx, including table content."""
    from docx import Document  # noqa: PLC0415

    document = Document(io.BytesIO(file_bytes))
    paragraphs: list[str] = []

    # Body paragraphs
    for para in document.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)

    # Table cells
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                paragraphs.append(row_text)

    return "\n".join(paragraphs)


def extract_text_from_docx(file_bytes: bytes) -> ExtractionResult:
    """Extract text from DOCX bytes using a tiered strategy."""
    strategies = [
        ("unstructured", _extract_docx_unstructured),
        ("python-docx", _extract_docx_python_docx),
    ]

    for method_name, extractor in strategies:
        try:
            text = extractor(file_bytes)
            if text.strip():
                logger.info("docx_extraction_success", method=method_name)
                return ExtractionResult(text=text, method=method_name)
            logger.debug("docx_extraction_empty", method=method_name)
        except Exception as exc:
            logger.warning("docx_extraction_failed", method=method_name, error=str(exc))

    warning = "Could not extract text from this DOCX file. It may be corrupted or unsupported."
    logger.warning("docx_extraction_all_failed")
    return ExtractionResult(text="", method="failed", warning=warning)


# ─── TXT extraction ───────────────────────────────────────────────────────


def extract_text_from_txt(file_bytes: bytes) -> ExtractionResult:
    """Decode TXT bytes, trying UTF-8 then latin-1."""
    try:
        text = file_bytes.decode("utf-8").strip()
    except UnicodeDecodeError:
        text = file_bytes.decode("latin-1", errors="ignore").strip()
    except Exception as exc:
        return ExtractionResult(text="", method="failed", warning=str(exc))

    return ExtractionResult(text=text, method="text")


# ─── Dispatcher ───────────────────────────────────────────────────────────


def extract_text(
    file_bytes: bytes,
    file_name: str,
    content_type: str | None = None,
    use_ocr: bool = False,
) -> ExtractionResult:
    """
    Entry point for text extraction.

    Dispatches to the appropriate extractor based on file extension and
    MIME type. Falls back gracefully when optional dependencies are absent.

    Args:
        file_bytes:   Raw bytes of the uploaded file.
        file_name:    Original file name (used for extension detection).
        content_type: MIME type from the upload (used as secondary signal).
        use_ocr:      Attempt OCR if all text-based PDF strategies fail.

    Returns:
        ExtractionResult with .text, .method, .word_count, .warning, .is_ocr
    """
    name_lower = (file_name or "").lower()

    if name_lower.endswith(".pdf") or content_type == "application/pdf":
        return extract_text_from_pdf(file_bytes, use_ocr=use_ocr)

    if name_lower.endswith(".docx") or content_type in {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/octet-stream",  # Some browsers send this for .docx
    }:
        # Re-check extension for octet-stream — don't mis-classify other binaries
        if name_lower.endswith(".docx") or content_type == (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            return extract_text_from_docx(file_bytes)

    if name_lower.endswith(".txt") or (content_type and content_type.startswith("text/")):
        return extract_text_from_txt(file_bytes)

    return ExtractionResult(
        text="",
        method="unsupported",
        warning=(
            f"Unsupported file type: '{file_name}'. "
            "Please upload a PDF, DOCX, or TXT file."
        ),
    )
